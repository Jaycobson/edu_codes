import pandas as pd
from docx import Document
from io import BytesIO

def export_to_csv(quiz_results: list) -> BytesIO:
    df = pd.DataFrame(quiz_results)
    # Reformat options for better CSV display if needed
    df['options'] = df['options'].apply(lambda x: "; ".join(x))
    
    csv_buffer = BytesIO()
    df.to_csv(csv_buffer, index=False)
    csv_buffer.seek(0)
    return csv_buffer

def export_to_docx(quiz_results: list) -> BytesIO:
    document = Document()
    document.add_heading('Quiz Results', 0)

    for i, result in enumerate(quiz_results):
        document.add_heading(f"Question {i+1}: {result['question']}", level=2)
        document.add_paragraph(f"Options: {', '.join(result['options'])}")
        document.add_paragraph(f"Your Answer: {result['chosen_answer']}")
        document.add_paragraph(f"Correct Answer: {result['correct_answer']}")
        
        if result['is_correct']:
            document.add_paragraph("Result: Correct")
        else:
            document.add_paragraph("Result: Incorrect")
        
        document.add_paragraph(f"Explanation: {result['explanation']}")
        document.add_page_break()

    doc_buffer = BytesIO()
    document.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer